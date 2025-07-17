from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from datetime import datetime, timedelta
import httpx
import logging
import os
import shutil
import uuid
import time
import hashlib
import docx2txt  # For basic DOCX processing
import PyPDF4   # For PDF processing
import base64
import io
import re  # Add missing re import
from docx import Document  # For advanced DOCX processing
from PIL import Image
import fitz  # PyMuPDF for better PDF processing
import magic  # For file type detection
import json
from pathlib import Path

# Import enhanced table processing utilities
from utils.table_parser import TableParser
from utils.response_formatter import ResponseFormatter
from utils.markdown_processor import MarkdownProcessor

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Folder untuk menyimpan file yang diunggah
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

class ChatRequest(BaseModel):
    message: str
    context: Optional[str] = None  # Untuk menyimpan konteks dari dokumen
    model_type: Optional[str] = "llama3:latest"  # Model selection
    conversation_history: Optional[List[Dict[str, str]]] = []  # Chat history untuk konteks

class DocumentResponse(BaseModel):
    document_id: str
    content: str
    filename: str

# Document Library System
import json
from pathlib import Path

# Document Metadata Storage
METADATA_FILE = os.path.join(UPLOAD_FOLDER, "documents_metadata.json")

class DocumentMetadata(BaseModel):
    document_id: str
    filename: str
    original_filename: str
    upload_date: str
    file_size: int
    file_type: str  # .pdf or .docx
    content_preview: str  # First 200 chars
    analysis_summary: Dict[str, Any]
    is_active: bool = False

class DocumentLibrary:
    """Manage document library with persistent storage"""
    
    def __init__(self):
        self.metadata_file = METADATA_FILE
        self.ensure_metadata_file()
    
    def ensure_metadata_file(self):
        """Ensure metadata file exists"""
        if not os.path.exists(self.metadata_file):
            self.save_metadata([])
    
    def load_metadata(self) -> List[DocumentMetadata]:
        """Load all document metadata"""
        try:
            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return [DocumentMetadata(**item) for item in data]
        except Exception as e:
            logger.error(f"Error loading metadata: {e}")
            return []
    
    def save_metadata(self, documents: List[DocumentMetadata]):
        """Save document metadata"""
        try:
            data = [doc.dict() for doc in documents]
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Error saving metadata: {e}")
    
    def add_document(self, metadata: DocumentMetadata):
        """Add new document to library"""
        documents = self.load_metadata()
        
        # Set all other documents as inactive
        for doc in documents:
            doc.is_active = False
        
        # Add new document as active
        metadata.is_active = True
        documents.append(metadata)
        
        self.save_metadata(documents)
        logger.info(f"Added document to library: {metadata.filename}")
    
    def get_all_documents(self) -> List[DocumentMetadata]:
        """Get all documents in library"""
        return self.load_metadata()
    
    def get_document(self, document_id: str) -> Optional[DocumentMetadata]:
        """Get specific document by ID"""
        documents = self.load_metadata()
        for doc in documents:
            if doc.document_id == document_id:
                return doc
        return None
    
    def set_active_document(self, document_id: str) -> bool:
        """Set document as active"""
        documents = self.load_metadata()
        found = False
        
        for doc in documents:
            if doc.document_id == document_id:
                doc.is_active = True
                found = True
            else:
                doc.is_active = False
        
        if found:
            self.save_metadata(documents)
            logger.info(f"Set active document: {document_id}")
        
        return found
    
    def get_active_document(self) -> Optional[DocumentMetadata]:
        """Get currently active document"""
        documents = self.load_metadata()
        for doc in documents:
            if doc.is_active:
                return doc
        return None
    
    def delete_document(self, document_id: str) -> bool:
        """Delete document from library and filesystem"""
        documents = self.load_metadata()
        doc_to_delete = None
        
        # Find document to delete
        for i, doc in enumerate(documents):
            if doc.document_id == document_id:
                doc_to_delete = doc
                documents.pop(i)
                break
        
        if doc_to_delete:
            # Delete physical file
            file_path = os.path.join(UPLOAD_FOLDER, doc_to_delete.filename)
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Deleted file: {file_path}")
            
            # Save updated metadata
            self.save_metadata(documents)
            logger.info(f"Deleted document from library: {doc_to_delete.original_filename}")
            return True
        
        return False

# Simplified Document Processing for Library System
class SimpleDocumentProcessor:
    """Simplified document processing for active document selection"""
    
    def enhance_context(self, query: str, full_context: str) -> str:
        """Speed-optimized context processing"""
        max_length = 6000  # REDUCED: 6KB context limit for speed (was 15000)
        
        if not full_context:
            return ""
        
        if len(full_context) <= max_length:
            return full_context.strip()
        
        # Smart truncation - keep most relevant parts based on query
        query_words = query.lower().split()
        
        # Try to find most relevant sections
        paragraphs = full_context.split('\n\n')
        relevant_paragraphs = []
        current_length = 0
        
        # First, add paragraphs that contain query words
        for para in paragraphs:
            if any(word in para.lower() for word in query_words):
                if current_length + len(para) <= max_length:
                    relevant_paragraphs.append(para)
                    current_length += len(para)
                else:
                    break
        
        # If we have space, add other paragraphs
        if current_length < max_length * 0.8:  # Use 80% threshold
            for para in paragraphs:
                if para not in relevant_paragraphs:
                    if current_length + len(para) <= max_length:
                        relevant_paragraphs.append(para)
                        current_length += len(para)
                    else:
                        break
        
        if relevant_paragraphs:
            return '\n\n'.join(relevant_paragraphs)
        else:
            # Fallback to simple truncation
            return full_context[:max_length] + "..."

class SimplePromptEngineer:
    """Simplified prompt engineering for document and general chat"""
    
    def create_document_prompt(self, query: str, context: str, conversation_history: List[Dict] = None) -> str:
        """Create speed-optimized document analysis prompt"""
        
        conversation_context = ""
        if conversation_history:
            last_exchanges = conversation_history[-1:] if conversation_history else []  # REDUCED: Only last 1 exchange for speed
            for msg in last_exchanges:
                conversation_context += f"{msg.get('sender', 'User')}: {msg.get('content', '')[:100]}...\n"
        
        # SPEED-OPTIMIZED prompt with clear instructions for fast response
        prompt = f"""INSTRUKSI CEPAT - JAWAB LANGSUNG:
Berikan jawaban singkat, padat, dan informatif dengan format markdown:
- Gunakan ## untuk judul utama
- **Bold** untuk poin penting (minimal 3-5 bold)
- - untuk daftar (gunakan banyak list)
- Jawaban maksimal 500-800 kata
- LANGSUNG ke inti masalah
- JANGAN menulis pembukaan panjang

DOKUMEN:
{context}

RIWAYAT: {conversation_context}

PERTANYAAN: {query}

JAWABAN LANGSUNG:"""
        
        return prompt
    
    def create_general_prompt(self, query: str, conversation_history: List[Dict] = None) -> str:
        """Create speed-optimized general conversation prompt"""
        
        conversation_context = ""
        if conversation_history:
            last_exchanges = conversation_history[-1:] if conversation_history else []  # REDUCED: Only last 1 exchange
            for msg in last_exchanges:
                conversation_context += f"{msg.get('sender', 'User')}: {msg.get('content', '')[:100]}...\n"
        
        # Check if it's likely a detailed question
        is_detailed_query = any(word in query.lower() for word in [
            'explain', 'jelaskan', 'how', 'bagaimana', 'apa itu', 'what is', 'describe', 'deskripsikan',
            'list', 'daftar', 'compare', 'bandingkan', 'analyze', 'analisis', 'cara', 'langkah', 'step'
        ]) or len(query) > 50  # REDUCED threshold
        
        if is_detailed_query:
            prompt = f"""INSTRUKSI CEPAT:
Jawab dengan format markdown yang rapi:
- ## untuk judul
- **Bold** untuk poin penting
- - untuk list
- Maksimal 400-600 kata
- Langsung ke inti

{conversation_context}

Pertanyaan: {query}
Jawaban singkat:"""
        else:
            prompt = f"""{conversation_context}

Pertanyaan: {query}
Jawaban:"""
        
        return prompt

# Simplified Performance Monitoring
class SimplePerformanceMonitor:
    """Speed-optimized performance monitoring"""
    
    def __init__(self):
        self.last_response_time = 0
        
    def add_response_time(self, response_time_ms: float):
        """Add a response time measurement"""
        self.last_response_time = response_time_ms
    
    def get_simple_status(self) -> str:
        """Get performance status - SPEED OPTIMIZED thresholds"""
        if self.last_response_time < 30000:     # 30 seconds
            return "excellent"
        elif self.last_response_time < 60000:   # 1 minute  
            return "good"
        elif self.last_response_time < 120000:  # 2 minutes
            return "acceptable"
        elif self.last_response_time < 180000:  # 3 minutes
            return "slow"
        else:
            return "timeout_risk"

# Simple Document Analysis Functions
async def analyze_docx_structure(file_path: str) -> dict:
    """Simple DOCX structure analysis"""
    try:
        doc = Document(file_path)
        
        return {
            "document_type": "DOCX",
            "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
            "tables": len(doc.tables),
            "text_length": sum(len(p.text) for p in doc.paragraphs),
        }
    except Exception as e:
        logger.error(f"Error analyzing DOCX: {e}")
        return {"error": str(e)}

async def analyze_pdf_structure(file_path: str) -> dict:
    """Simple PDF structure analysis"""
    try:
        pdf_document = fitz.open(file_path)
        
        text_length = 0
        images = 0
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text_length += len(page.get_text())
            images += len(page.get_images())
        
        pdf_document.close()
        
        return {
            "document_type": "PDF",
            "pages": len(pdf_document),
            "images": images,
            "text_length": text_length,
        }
    except Exception as e:
        logger.error(f"Error analyzing PDF: {e}")
        return {"error": str(e)}

# Simplified Response Processing Classes
class ResponseCache:
    """Simple in-memory cache for responses"""
    
    def __init__(self, max_size: int = 100, ttl_minutes: int = 30):
        self.cache: Dict[str, Dict] = {}
        self.max_size = max_size
        self.ttl = timedelta(minutes=ttl_minutes)
    
    def _generate_key(self, prompt: str, context: str = None) -> str:
        """Generate cache key"""
        content = f"{prompt}:{context or ''}"
        return hashlib.md5(content.encode()).hexdigest()
    
    def get(self, prompt: str, context: str = None) -> Optional[str]:
        """Get cached response"""
        key = self._generate_key(prompt, context)
        
        if key in self.cache:
            entry = self.cache[key]
            if datetime.now() - entry['timestamp'] < self.ttl:
                return entry['response']
            else:
                del self.cache[key]
        
        return None
    
    def set(self, prompt: str, response: str, context: str = None):
        """Cache a response"""
        key = self._generate_key(prompt, context)
        
        if len(self.cache) >= self.max_size:
            oldest_key = min(self.cache.keys(), 
                           key=lambda k: self.cache[k]['timestamp'])
            del self.cache[oldest_key]
        
        self.cache[key] = {
            'response': response,
            'timestamp': datetime.now()
        }
    
    def clear(self):
        """Clear cache"""
        self.cache.clear()

class ResponseMonitor:
    """Simple response quality monitor"""
    
    def clean_response(self, text: str) -> str:
        """Basic response cleaning"""
        if not text or text.strip() == "":
            return "Maaf, tidak ada jawaban yang dapat saya berikan."
        
        cleaned_text = text.strip()
        
        # Limit length if extremely long - INCREASED for long responses
        if len(cleaned_text) > 80000:  # INCREASED: ~15,000 words support (was 2000)
            cleaned_text = cleaned_text[:80000] + "..."
        
        return cleaned_text

# AI Model Configuration System
@dataclass
class ModelConfig:
    model_name: str
    temperature: float
    top_p: float
    top_k: int
    num_ctx: int
    num_predict: int
    repeat_penalty: float
    num_thread: int
    stop_tokens: list

class AIModelOptimizer:
    """Ultra-fast AI Model optimizer for gaming laptop - SPEED OPTIMIZED"""
    
    def __init__(self):
        # BALANCED configuration for quality + speed optimization
        self.config = ModelConfig(
            model_name="llama3:8b",
            temperature=0.6,        # FURTHER REDUCED: Even faster, still good quality  
            top_p=0.8,             # REDUCED: More focused for speed
            top_k=25,              # FURTHER REDUCED: Lower for faster generation
            num_ctx=4096,          # Keep current - good balance
            num_predict=2500,      # REDUCED: Target ~1500-2000 words for speed (was 3000)
            repeat_penalty=1.2,    # Slightly higher to prevent loops and reduce iteration
            num_thread=-1,         # Use all CPU threads
            stop_tokens=["Human:", "Assistant:", "PERTANYAAN:", "User:", "\n\nHuman", "\n\nUser"]
        )
    
    def get_config(self) -> ModelConfig:
        return self.config
    
    def get_optimized_payload(self, prompt: str, use_streaming: bool = True) -> Dict[str, Any]:
        config = self.get_config()
        
        return {
            "model": config.model_name,
            "prompt": prompt,
            "stream": use_streaming,
            "options": {
                "temperature": config.temperature,
                "top_p": config.top_p,
                "top_k": config.top_k,
                "num_ctx": config.num_ctx,
                "num_predict": config.num_predict,
                "num_thread": config.num_thread,
                "repeat_penalty": config.repeat_penalty,
                "repeat_last_n": 30,   # REDUCED: Smaller repeat window for speed
                "stop": config.stop_tokens,
                
                # SPEED-OPTIMIZED configuration for sub-60s responses
                "num_gpu": -1,         # Auto-detect GPU 
                "low_vram": False,     # Keep disabled for 22GB RAM
                "f16_kv": True,        # ENABLED: Use half precision for speed
                "num_batch": 512,      # REDUCED: Smaller batch for faster processing
                "numa": False,         # Keep disabled
                "mlock": True,         # Lock model in memory
                "use_mmap": True,      # Enable memory mapping
                "seed": -1,            # Random seed
                
                # AGGRESSIVE speed optimizations
                "penalize_newline": False,
                "presence_penalty": 0.0,
                "frequency_penalty": 0.0,
                "tfs_z": 1.0,
                "typical_p": 1.0,
                "min_p": 0.0,
                
                # Additional speed tweaks
                "rope_freq_base": 10000,
                "rope_freq_scale": 1.0,
                "num_keep": 24,        # Keep first 24 tokens for consistency
                "num_gqa": -1
            }
        }

# Initialize optimizer
ai_optimizer = AIModelOptimizer()

# Initialize document library
document_library = DocumentLibrary()

# Initialize simple processors
doc_processor = SimpleDocumentProcessor()
prompt_engineer = SimplePromptEngineer()

# Initialize performance monitoring and caching
performance_monitor = SimplePerformanceMonitor()
response_cache = ResponseCache(max_size=200, ttl_minutes=60)
response_monitor = ResponseMonitor()

# Initialize enhanced table processing
table_parser = TableParser()
response_formatter = ResponseFormatter()
markdown_processor = MarkdownProcessor()

# Streaming Response Handler for Long Responses
class StreamingResponseHandler:
    """Handle streaming responses to prevent frontend timeout"""
    
    async def process_streaming_response(self, response, timeout_seconds: float) -> str:
        """Process streaming response with timeout protection - SPEED OPTIMIZED"""
        full_response = ""
        start_time = time.time()
        last_chunk_time = time.time()
        chunk_timeout = 120.0  # ULTRA-EXTENDED: 120 seconds between chunks for very complex processing (was 60)
        progress_threshold = 50  # Log every 50 characters for better monitoring
        
        try:
            async for line in response.aiter_lines():
                current_time = time.time()
                
                # Check overall timeout
                if (current_time - start_time) > timeout_seconds:
                    logger.warning(f"⏱️ Overall timeout reached: {timeout_seconds}s")
                    break
                
                # Check chunk timeout (faster detection)
                if (current_time - last_chunk_time) > chunk_timeout:
                    logger.warning(f"⏱️ Chunk timeout reached: {chunk_timeout}s")
                    break
                
                if line:
                    try:
                        chunk_data = json.loads(line)
                        if "response" in chunk_data:
                            chunk_text = chunk_data["response"]
                            full_response += chunk_text
                            last_chunk_time = current_time
                            
                            # More frequent progress logging for speed monitoring
                            if len(full_response) % progress_threshold == 0:
                                elapsed = current_time - start_time
                                chars_per_sec = len(full_response) / elapsed if elapsed > 0 else 0
                                logger.info(f"📝 Streaming: {len(full_response)} chars ({chars_per_sec:.1f} chars/s)")
                        
                        # Check if done
                        if chunk_data.get("done", False):
                            elapsed = current_time - start_time
                            chars_per_sec = len(full_response) / elapsed if elapsed > 0 else 0
                            logger.info(f"✅ Streaming completed: {len(full_response)} chars in {elapsed:.1f}s ({chars_per_sec:.1f} chars/s)")
                            break
                            
                    except json.JSONDecodeError:
                        continue
                        
        except Exception as e:
            logger.error(f"❌ Streaming error: {e}")
            
        return full_response
    
    async def fallback_to_non_streaming(self, client, request_payload: dict, timeout_seconds: float) -> str:
        """Fallback to non-streaming if streaming fails"""
        logger.info("🔄 Falling back to non-streaming mode")
        
        # Modify payload for non-streaming
        request_payload["stream"] = False
        
        response = await client.post(
            "http://localhost:11434/api/generate",
            json=request_payload,
            timeout=timeout_seconds
        )
        
        if response.status_code == 200:
            result = response.json()
            return result.get("response", "")
        else:
            raise HTTPException(status_code=500, detail=f"Ollama API error: {response.status_code}")

# Initialize streaming handler
streaming_handler = StreamingResponseHandler()

@app.post("/api/chat")
async def chat_with_llama(request: ChatRequest):
    # Start timing
    start_time = time.time()
    request_start = datetime.now()
    
    try:
        logger.info(f"🚀 [TIMING] Chat request started at {request_start.strftime('%H:%M:%S.%f')[:-3]}")
        logger.info(f"📩 Request: {request.message[:100]}...")
        logger.info(f"📊 Context size: {len(request.context) if request.context else 0} chars")
        logger.info(f"💬 History: {len(request.conversation_history) if request.conversation_history else 0} messages")
        
        # Check for very long context that might cause timeouts
        context_warning_threshold = 3000  # REDUCED: Warn at 3KB instead of 5KB
        if request.context and len(request.context) > context_warning_threshold:
            logger.warning(f"⚠️ Large document context detected: {len(request.context)} characters")
            logger.warning("Sistem dioptimalkan untuk respons cepat. Pertanyaan spesifik akan lebih cepat.")
        
        # Check for very long messages
        if len(request.message) > 200:  # REDUCED: Warn at 200 chars instead of 500
            logger.warning(f"⚠️ Long message detected: {len(request.message)} characters")
            logger.warning("Pertanyaan yang lebih singkat dan spesifik akan mendapat respons lebih cepat.")

        # Enhanced prompt engineering with conversation history and active document
        prompt_start = time.time()
        
        # Check for active document if no context provided
        final_context = request.context
        active_doc_info = ""
        
        # CRITICAL DEBUG: Log current active document
        current_active = document_library.get_active_document()
        if current_active:
            logger.info(f"🔍 [DEBUG] Current active document: {current_active.original_filename} (ID: {current_active.document_id})")
        else:
            logger.info("🔍 [DEBUG] No active document currently set")
        
        if not final_context:
            # Check for active document in library
            active_doc = document_library.get_active_document()
            if active_doc:
                # Load content from active document
                file_path = os.path.join(UPLOAD_FOLDER, active_doc.filename)
                if os.path.exists(file_path):
                    try:
                        final_context = await extract_text_from_document(file_path)
                        active_doc_info = f"📄 Using active document: {active_doc.original_filename}"
                        logger.info(f"Using active document: {active_doc.original_filename}")
                    except Exception as e:
                        logger.warning(f"Failed to load active document: {e}")
        
        if final_context:
            # SPEED: Limit document context size aggressively for fast responses
            max_context = 8000  # REDUCED: Max 8KB context for speed (was 25000)
            if len(final_context) > max_context:
                logger.warning(f"⚠️ Document truncated for speed: {len(final_context)} → {max_context} chars")
                final_context = final_context[:max_context] + "..."
            
            # Process document with speed-optimized context
            try:
                enhanced_context = doc_processor.enhance_context(request.message, final_context)
                optimized_prompt = prompt_engineer.create_document_prompt(
                    request.message, 
                    enhanced_context,
                    request.conversation_history or []
                )
                logger.info(f"Using speed-optimized document context with {len(enhanced_context)} characters")
            except Exception as e:
                logger.error(f"Document processing failed: {e}")
                # Fallback to very simple context for speed
                enhanced_context = final_context[:1000] + "..." if len(final_context) > 1000 else final_context
                optimized_prompt = prompt_engineer.create_document_prompt(
                    request.message, 
                    enhanced_context,
                    request.conversation_history or []
                )
                logger.info(f"Using speed fallback context with {len(enhanced_context)} characters")
        else:
            # Enhanced general conversation with history
            optimized_prompt = prompt_engineer.create_general_prompt(
                request.message,
                request.conversation_history or []
            )
            logger.info("Using speed-optimized general conversation prompt")
        
        prompt_time = (time.time() - prompt_start) * 1000
        logger.info(f"⚡ [TIMING] Prompt engineering: {prompt_time:.2f}ms")
        
        # Check cache first
        cache_start = time.time()
        cached_response = response_cache.get(request.message, context=final_context)
        cache_time = (time.time() - cache_start) * 1000
        logger.info(f"⚡ [TIMING] Cache check: {cache_time:.2f}ms")
        
        # CRITICAL DEBUG: Log cache hit/miss with context info
        if cached_response:
            logger.info(f"🎯 [CACHE] HIT - Using cached response for message: {request.message[:50]}...")
            if final_context:
                logger.info(f"🎯 [CACHE] Context length: {len(final_context)} chars")
        else:
            logger.info(f"❌ [CACHE] MISS - Will query Ollama for message: {request.message[:50]}...")
        
        if cached_response:
            total_time = (time.time() - start_time) * 1000
            logger.info(f"🎯 [TIMING] CACHE HIT! Total response time: {total_time:.2f}ms")
            
            # Add enhanced document info to cached response if available
            response_data = {
                "response": cached_response,
                "status": "complete",
                "cached": True,
                "timing": {
                    "total_ms": total_time,
                    "source": "cache"
                },
                "chat_context": {
                    "has_document_context": bool(final_context),
                    "is_document_chat": bool(final_context),
                    "document_source": "active_document" if not request.context else "uploaded_document"
                }
            }
            
            # ✨ Apply table formatting to cached responses too
            try:
                cached_formatted = response_formatter.format_response(cached_response)
                cached_frontend = response_formatter.format_for_frontend(cached_formatted)
                
                response_data["enhanced_formatting"] = {
                    "has_tables": cached_formatted.has_tables,
                    "has_enhanced_content": cached_formatted.has_enhanced_content,
                    "table_count": cached_formatted.metadata.get('table_count', 0),
                    "formatting_applied": cached_formatted.formatting_applied,
                    "frontend_data": cached_frontend
                }
                
                if cached_formatted.has_tables:
                    response_data["table_metadata"] = {
                        "tables": [
                            {
                                "id": table.id,
                                "headers": table.headers,
                                "row_count": len(table.rows),
                                "column_count": len(table.headers),
                                "column_types": table.column_types
                            }
                            for table in cached_formatted.tables
                        ],
                        "rendering_hints": cached_frontend.get("rendering_hints", {})
                    }
                    
            except Exception as cache_format_error:
                logger.warning(f"Cache formatting failed: {cache_format_error}")
                response_data["enhanced_formatting"] = {
                    "has_tables": False,
                    "has_enhanced_content": False,
                    "table_count": 0,
                    "formatting_applied": ["basic"],
                    "frontend_data": None
                }
            
            # Add enhanced document context for cached responses
            if active_doc_info or final_context:
                active_doc = document_library.get_active_document()
                response_data["document_context"] = {
                    "display_name": active_doc.original_filename if active_doc else "Uploaded Document",
                    "file_type": active_doc.file_type if active_doc else "unknown",
                    "document_id": active_doc.document_id if active_doc else None,
                    "content_length": len(final_context) if final_context else 0,
                    "context_info": active_doc_info if active_doc_info else "📄 Analyzing uploaded document"
                }
            
            return JSONResponse(response_data)
        
        # Prepare request payload
        payload_start = time.time()
        
        # Set timeout for comprehensive analysis - ULTRA-EXTENDED FOR COMPLEX DOCUMENT ANALYSIS
        timeout_seconds = 1500.0  # ULTRA-EXTENDED: 25 minutes for complex document analysis (was 720)
        
        async with httpx.AsyncClient(timeout=timeout_seconds) as client:
            # SPEED: Always use streaming for responsiveness, but with smaller targets
            use_streaming = True
            
            # Get optimized payload with streaming option
            request_payload = ai_optimizer.get_optimized_payload(optimized_prompt, use_streaming)
            
            payload_time = (time.time() - payload_start) * 1000
            logger.info(f"⚡ [TIMING] Payload preparation: {payload_time:.2f}ms")
            
            logger.info(f"🚀 EXTENDED-TIMEOUT llama3:8b configuration")
            logger.info(f"Context length: {request_payload['options']['num_ctx']}")
            logger.info(f"Max tokens: {request_payload['options']['num_predict']}")
            logger.info(f"Temperature: {request_payload['options']['temperature']}")
            logger.info(f"Batch size: {request_payload['options'].get('num_batch', 'default')}")
            logger.info(f"Timeout: {timeout_seconds}s (25 MINUTES FOR COMPLEX DOCUMENT ANALYSIS)")
            logger.info(f"Prompt length: {len(optimized_prompt)} characters")
            logger.info(f"🌊 Streaming: ALWAYS ENABLED for responsiveness")
            
            # Make request to Ollama with streaming support
            ollama_start = time.time()
            logger.info(f"🤖 [TIMING] Sending streaming request to Ollama at {datetime.now().strftime('%H:%M:%S.%f')[:-3]}")
            
            try:
                # Always use streaming for responsiveness
                response = await client.post(
                    "http://localhost:11434/api/generate",
                    json=request_payload,
                    timeout=timeout_seconds
                )
                
                if response.status_code == 200:
                    # Process streaming response with shorter timeout
                    logger.info("🌊 Processing streaming response (SPEED MODE)...")
                    ai_response = await streaming_handler.process_streaming_response(
                        response, 
                        timeout_seconds * 0.8  # Use 80% of timeout for safety
                    )
                    
                    if not ai_response or len(ai_response.strip()) < 10:
                        logger.warning("⚠️ Streaming response too short, using emergency fallback")
                        # Emergency short response fallback
                        emergency_payload = ai_optimizer.get_optimized_payload(optimized_prompt, False)
                        emergency_payload["options"]["num_predict"] = 800  # Very short response
                        emergency_payload["options"]["temperature"] = 0.3  # Very focused
                        
                        response = await client.post(
                            "http://localhost:11434/api/generate",
                            json=emergency_payload,
                            timeout=600.0  # 10 minutes emergency timeout for very complex queries
                        )
                        
                        if response.status_code == 200:
                            result = response.json()
                            ai_response = result.get("response", "") + "\n\n*[Respons dipercepat untuk menghindari timeout]*"
                        else:
                            raise HTTPException(status_code=500, detail=f"Emergency fallback failed: {response.status_code}")
                else:
                    raise HTTPException(status_code=500, detail=f"Ollama streaming error: {response.status_code}")
                        
            except httpx.TimeoutException:
                logger.error("⏱️ Request timeout, attempting ULTRA-FAST emergency fallback...")
                try:
                    # Ultra-fast emergency response
                    ultra_fast_payload = ai_optimizer.get_optimized_payload(
                        f"Berikan jawaban singkat dan padat untuk: {request.message}", False
                    )
                    ultra_fast_payload["options"]["num_predict"] = 300  # Very short
                    ultra_fast_payload["options"]["temperature"] = 0.1  # Very focused
                    ultra_fast_payload["options"]["top_k"] = 10        # Very selective
                    
                    response = await client.post(
                        "http://localhost:11434/api/generate",
                        json=ultra_fast_payload,
                        timeout=30.0  # 30 seconds ultra-fast timeout
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        ai_response = result.get("response", "") + "\n\n*[Mode respons cepat karena timeout - silakan coba lagi untuk jawaban lebih detail]*"
                        logger.info("✅ Ultra-fast emergency fallback successful")
                    else:
                        raise
                        
                except Exception as emergency_error:
                    logger.error(f"❌ Ultra-fast emergency fallback failed: {emergency_error}")
                    raise HTTPException(
                        status_code=504, 
                        detail="🚀 Sistem dioptimalkan untuk respons cepat. Silakan coba pertanyaan yang lebih spesifik atau singkat."
                    )
            
            ollama_end = time.time()
            ollama_time = (ollama_end - ollama_start) * 1000
            logger.info(f"🤖 [TIMING] Ollama response received: {ollama_time:.2f}ms ({ollama_time/1000:.1f}s)")
            
            # Handle response processing with timing
            processing_start = time.time()
            
            # Debug logging
            logger.info(f"AI response length: {len(ai_response)}")
            logger.info(f"AI response preview: {ai_response[:100]}...")
            
            # VALIDASI: Jika response terlalu pendek, ini kemungkinan error
            if len(ai_response.strip()) < 5:  # Kurang dari 5 karakter
                logger.warning(f"⚠️ Response terlalu pendek ({len(ai_response)} chars): '{ai_response}'")
                logger.warning("Kemungkinan stop tokens terlalu agresif atau model error")
                ai_response = f"Maaf, sistem menghasilkan respons yang tidak lengkap ('{ai_response.strip()}'). Silakan coba lagi dengan pertanyaan yang berbeda."
            
            # If response is empty, provide a fallback
            if not ai_response or ai_response.strip() == "":
                logger.warning("Empty response from Ollama, using fallback")
                ai_response = "Maaf, saya tidak dapat memberikan jawaban saat ini. Silakan coba lagi dengan pertanyaan yang berbeda."
            
            # Clean the response
            cleaned_response = response_monitor.clean_response(ai_response)
            
            # ✨ NEW: Enhanced table processing
            try:
                formatted_response = response_formatter.format_response(cleaned_response)
                frontend_formatted = response_formatter.format_for_frontend(formatted_response)
                
                logger.info(f"📊 Table processing: {formatted_response.metadata.get('table_count', 0)} tables detected")
                if formatted_response.has_tables:
                    logger.info(f"🎨 Enhanced formatting applied: {formatted_response.formatting_applied}")
                    
                    # ✨ CRITICAL: Remove markdown tables from content when enhanced tables are available
                    cleaned_response = response_formatter.remove_markdown_tables_from_text(cleaned_response)
                    logger.info("✂️ Markdown tables removed from response content")
                        
            except Exception as format_error:
                logger.warning(f"Table formatting failed, using fallback: {format_error}")
                # Fallback to basic response
                formatted_response = None
                frontend_formatted = None
            
            # ✨ NEW: Enhanced markdown processing for rich text formatting
            markdown_metadata = None
            try:
                # IMPROVED: More sensitive markdown detection for better formatting
                markdown_indicators = [
                    # Heading indicators (more sensitive)
                    '#' in cleaned_response and len(cleaned_response.split('#')) >= 2,  # At least 1 heading
                    re.search(r'^#{1,6}\s+.+', cleaned_response, re.MULTILINE),  # Proper heading format
                    
                    # Structure indicators
                    '```' in cleaned_response,  # Code blocks
                    cleaned_response.count('**') >= 2,  # At least 1 bold pair
                    
                    # List indicators (more sensitive)
                    re.search(r'^\s*[-*+]\s+.+', cleaned_response, re.MULTILINE),  # Any bullet list
                    re.search(r'^\s*\d+\.\s+.+', cleaned_response, re.MULTILINE),  # Any numbered list
                    
                    # Blockquote indicators
                    re.search(r'^>\s+.+', cleaned_response, re.MULTILINE),  # Any blockquote
                    
                    # Multi-line structure (indicates formatted content)
                    len(cleaned_response.split('\n')) > 3 and any([
                        '#' in cleaned_response,
                        '**' in cleaned_response,
                        '- ' in cleaned_response,
                        '1. ' in cleaned_response
                    ])
                ]
                
                has_markdown_content = any(markdown_indicators)
                
                # DEBUG: Log markdown detection details
                logger.info(f"🔍 [MARKDOWN DEBUG] Response length: {len(cleaned_response)}")
                logger.info(f"🔍 [MARKDOWN DEBUG] Has markdown indicators: {has_markdown_content}")
                logger.info(f"🔍 [MARKDOWN DEBUG] Indicators: {markdown_indicators}")
                logger.info(f"🔍 [MARKDOWN DEBUG] Response preview: {cleaned_response[:200]}...")
                
                # Apply markdown processing if there's ANY markdown content AND reasonable length
                if has_markdown_content and len(cleaned_response) > 50:  # Lower threshold
                    logger.info("📝 Markdown content detected, processing...")
                    
                    # Always process markdown when detected
                    enhanced_markdown = cleaned_response
                    
                    # Process the markdown
                    processed_markdown = markdown_processor.process_markdown(enhanced_markdown)
                    markdown_metadata = {
                        "is_markdown": True,
                        "has_headings": processed_markdown.metadata.has_headings,
                        "has_lists": processed_markdown.metadata.has_lists,
                        "has_emphasis": processed_markdown.metadata.has_emphasis,
                        "has_code_blocks": processed_markdown.metadata.has_code_blocks,
                        "has_links": processed_markdown.metadata.has_links,
                        "has_tables": processed_markdown.metadata.has_tables,
                        "has_blockquotes": processed_markdown.metadata.has_blockquotes,
                        "word_count": processed_markdown.metadata.word_count,
                        "formatting_applied": processed_markdown.formatting_applied
                    }
                    
                    # Use the enhanced markdown content
                    cleaned_response = processed_markdown.raw_markdown
                    logger.info(f"📝 Markdown processing: {len(processed_markdown.formatting_applied)} formats applied")
                else:
                    logger.info("📝 Skipping markdown processing - insufficient markdown content detected")
                    
                    # Even without full processing, check if we have basic markdown for frontend
                    has_basic_markdown = any([
                        bool(re.search(r'^#{1,6}\s+.+', cleaned_response, re.MULTILINE)),  # Has headings
                        '```' in cleaned_response,  # Has code blocks
                        cleaned_response.count('**') >= 2,  # Has bold text
                        bool(re.search(r'^\s*[-*+]\s+.+', cleaned_response, re.MULTILINE)),  # Has lists
                        bool(re.search(r'^\s*\d+\.\s+.+', cleaned_response, re.MULTILINE)),  # Has numbered lists
                        bool(re.search(r'^>\s+.+', cleaned_response, re.MULTILINE))  # Has blockquotes
                    ])
                    
                    if has_basic_markdown:
                        logger.info("📝 Basic markdown elements detected, preparing frontend metadata")
                        markdown_metadata = {
                            "is_markdown": True,
                            "has_headings": bool(re.search(r'^#{1,6}\s+.+', cleaned_response, re.MULTILINE)),
                            "has_lists": bool(re.search(r'^\s*[-*+\d]\s+.+', cleaned_response, re.MULTILINE)),
                            "has_emphasis": cleaned_response.count('**') >= 2 or cleaned_response.count('*') >= 2,
                            "has_code_blocks": '```' in cleaned_response,
                            "has_links": bool(re.search(r'\[.*?\]\(.*?\)', cleaned_response)),
                            "has_tables": bool(re.search(r'\|.*\|', cleaned_response)),
                            "has_blockquotes": bool(re.search(r'^>\s+.+', cleaned_response, re.MULTILINE)),
                            "word_count": len(cleaned_response.split()),
                            "formatting_applied": []
                        }
                        
            except Exception as markdown_error:
                logger.warning(f"Markdown processing failed, using fallback: {markdown_error}")
                markdown_metadata = None
            
            # Cache the response
            response_cache.set(request.message, cleaned_response, final_context)
            
            # Monitor performance
            performance_monitor.add_response_time(ollama_time)
            
            processing_time = (time.time() - processing_start) * 1000
            total_time = (time.time() - start_time) * 1000
            
            # Add simple performance monitoring
            perf_status = performance_monitor.get_simple_status()
            
            logger.info(f"⚡ [TIMING] Response processing: {processing_time:.2f}ms")
            logger.info(f"🎯 [TIMING] TOTAL REQUEST TIME: {total_time:.2f}ms ({total_time/1000:.2f}s)")
            logger.info(f"📊 [PERFORMANCE] Status: {perf_status}")
            
            # Performance warnings and recommendations
            if perf_status == 'timeout_risk':
                logger.error("🚨 [PERFORMANCE] TIMEOUT RISK! Response took too long!")
            elif perf_status == 'slow':
                logger.warning("🐌 [PERFORMANCE] Slow response detected!")
            elif perf_status == 'acceptable':
                logger.info("⚡ [PERFORMANCE] Acceptable speed")
            elif perf_status in ['good', 'excellent']:
                logger.info("🚀 [PERFORMANCE] Good speed achieved!")
            
            # Add speed optimization suggestions based on performance
            speed_suggestion = ""
            if perf_status in ['slow', 'timeout_risk']:
                if final_context:
                    speed_suggestion = "\n\n💡 *Tip: Untuk respons lebih cepat, coba tanyakan tentang bagian spesifik dokumen*"
                else:
                    speed_suggestion = "\n\n💡 *Tip: Pertanyaan yang lebih singkat akan mendapat respons lebih cepat*"
            
            # Add speed suggestion to response if needed
            if speed_suggestion and perf_status in ['slow', 'timeout_risk']:
                cleaned_response += speed_suggestion
            
            logger.info(f"Final cleaned response length: {len(cleaned_response)} characters")
            
            # Prepare response data with enhanced document context and streaming info
            response_data = {
                "response": cleaned_response,
                "status": "complete",
                "streaming_used": use_streaming,  # Inform frontend about streaming
                "word_count": len(cleaned_response.split()),  # Add word count
                "timing": {
                    "total_ms": total_time,
                    "ollama_ms": ollama_time,
                    "processing_ms": processing_time,
                    "prompt_ms": prompt_time,
                    "cache_ms": cache_time,
                    "performance_status": perf_status,
                    "breakdown": {
                        "prompt_engineering": prompt_time,
                        "cache_check": cache_time,
                        "payload_prep": payload_time,
                        "ollama_request": ollama_time,
                        "response_processing": processing_time
                    }
                },
                "chat_context": {
                    "has_document_context": bool(final_context),
                    "is_document_chat": bool(final_context),
                    "document_source": "active_document" if not request.context else "uploaded_document"
                }
            }
            
            # ✨ Add enhanced table formatting data
            if formatted_response and frontend_formatted:
                response_data["enhanced_formatting"] = True
                response_data["table_metadata"] = [
                    {
                        "type": "table",
                        "headers": table.headers,
                        "rows": table.rows,
                        "column_types": table.column_types,
                        "title": getattr(table, 'title', None)
                    }
                    for table in formatted_response.tables
                ]
                
                logger.info(f"✅ Enhanced table data prepared: {len(response_data['table_metadata'])} tables")
            else:
                response_data["enhanced_formatting"] = False
                response_data["table_metadata"] = []
            
            # ✨ Add enhanced markdown formatting data
            if markdown_metadata:
                response_data["markdown_formatting"] = True
                response_data["markdown_metadata"] = markdown_metadata
                logger.info(f"✅ Markdown data prepared: {len(markdown_metadata.get('formatting_applied', []))} formats")
                logger.info(f"🔍 [FRONTEND DEBUG] Sending markdown_metadata: {markdown_metadata}")
            else:
                response_data["markdown_formatting"] = False
                response_data["markdown_metadata"] = {}
                logger.info("❌ [FRONTEND DEBUG] No markdown metadata - sending empty")
            
            # DEBUG: Log final response data being sent to frontend
            logger.info(f"🔍 [FRONTEND DEBUG] Final response keys: {list(response_data.keys())}")
            logger.info(f"🔍 [FRONTEND DEBUG] markdown_formatting: {response_data.get('markdown_formatting', False)}")
            logger.info(f"🔍 [FRONTEND DEBUG] Response content preview: {cleaned_response[:100]}...")
            
            # Add enhanced document info for chat display
            if active_doc_info or final_context:
                active_doc = document_library.get_active_document()
                response_data["document_context"] = {
                    "display_name": active_doc.original_filename if active_doc else "Uploaded Document",
                    "file_type": active_doc.file_type if active_doc else "unknown",
                    "document_id": active_doc.document_id if active_doc else None,
                    "content_length": len(final_context) if final_context else 0,
                    "context_info": active_doc_info if active_doc_info else "📄 Analyzing uploaded document"
                }
            
            return JSONResponse(response_data)
                
    except httpx.ConnectError:
        total_time = (time.time() - start_time) * 1000
        error_msg = "Failed to connect to Ollama. Please ensure Ollama is running on http://localhost:11434"
        logger.error(f"❌ [TIMING] Connection error after {total_time:.2f}ms: {error_msg}")
        raise HTTPException(status_code=503, detail=error_msg)
    except httpx.TimeoutException:
        total_time = (time.time() - start_time) * 1000
        
        # Speed-optimized timeout handling with clear user guidance
        if total_time > 120000:  # More than 2 minutes
            timeout_msg = f"⚡ Sistem dioptimalkan untuk respons cepat (target <60 detik). "
            timeout_msg += f"Request ini membutuhkan {total_time/1000:.1f} detik. "
            
            if request.context or final_context:
                timeout_msg += "💡 Untuk dokumen kompleks, coba tanyakan: 'Apa poin utama?' atau 'Ringkas dalam 5 poin'."
            else:
                timeout_msg += "💡 Coba dengan pertanyaan yang lebih spesifik dan singkat."
        else:
            timeout_msg = f"⏱️ Timeout setelah {total_time/1000:.1f}s. "
            if request.context:
                timeout_msg += "Untuk analisis dokumen, tanyakan tentang bagian spesifik saja."
            else:
                timeout_msg += "Coba dengan pertanyaan yang lebih fokus."
        
        logger.error(f"⏱️ [TIMING] Timeout after {total_time:.2f}ms")
        raise HTTPException(status_code=504, detail=timeout_msg)
    except Exception as e:
        total_time = (time.time() - start_time) * 1000
        error_msg = f"Unexpected error after {total_time:.2f}ms: {str(e)}"
        logger.error(f"❌ [TIMING] {error_msg}")
        raise HTTPException(status_code=500, detail=error_msg)

async def extract_text_from_document(file_path: str) -> str:
    """Ekstrak teks dan deskripsi konten dari file dokumen (PDF atau DOCX) dengan support gambar."""
    try:
        # Log untuk debugging
        logger.info(f"Extracting text from document: {file_path}")
        
        # Deteksi tipe file dengan python-magic untuk akurasi lebih baik
        try:
            file_type = magic.from_file(file_path, mime=True)
            logger.info(f"Detected MIME type: {file_type}")
        except:
            # Fallback ke ekstensi file
            file_extension = os.path.splitext(file_path)[1].lower()
            logger.info(f"Fallback to file extension: {file_extension}")
            file_type = None
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf' or (file_type and 'pdf' in file_type):
            return await extract_text_from_pdf_advanced(file_path)
        elif file_extension == '.docx' or (file_type and 'wordprocessingml' in file_type):
            return await extract_text_from_docx_advanced(file_path)
        else:
            logger.error(f"Format file tidak didukung: {file_extension}")
            return f"Format file tidak didukung: {file_extension}"
    except Exception as e:
        logger.error(f"Error saat mengekstrak teks dari dokumen: {e}", exc_info=True)
        return f"Error saat mengekstrak teks: {str(e)}"

async def extract_text_from_docx_advanced(file_path: str) -> str:
    """Ekstrak teks dari file DOCX dengan support lengkap untuk semua elemen."""
    try:
        logger.info("Starting advanced DOCX extraction...")
        
        # Gunakan python-docx untuk ekstraksi yang lebih lengkap
        doc = Document(file_path)
        
        extracted_content = []
        image_count = 0
        table_count = 0
        
        # Header
        extracted_content.append("=== DOKUMEN WORD ===\n")
        
        # Ekstrak paragraf dengan formatting info
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # Deteksi style/heading
                if paragraph.style.name.startswith('Heading'):
                    extracted_content.append(f"\n## {text}")
                elif paragraph.style.name == 'Title':
                    extracted_content.append(f"\n# {text}")
                else:
                    extracted_content.append(text)
        
        # Ekstrak tabel
        for table_idx, table in enumerate(doc.tables):
            table_count += 1
            extracted_content.append(f"\n[TABEL {table_count}]")
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_data.append(cell_text)
                
                if any(row_data):  # Hanya tambahkan jika ada data
                    extracted_content.append(" | ".join(row_data))
        
        # Ekstrak gambar dan beri deskripsi
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                extracted_content.append(f"\n[GAMBAR {image_count}] - Gambar ditemukan dalam dokumen")
        
        # Ekstrak properties dokumen
        if doc.core_properties.title:
            extracted_content.insert(1, f"Judul: {doc.core_properties.title}")
        if doc.core_properties.author:
            extracted_content.insert(2, f"Penulis: {doc.core_properties.author}")
        if doc.core_properties.subject:
            extracted_content.insert(3, f"Subjek: {doc.core_properties.subject}")
        
        # Tambahkan ringkasan konten
        summary = f"\n=== RINGKASAN DOKUMEN ===\n"
        summary += f"- Total paragraf: {len([p for p in doc.paragraphs if p.text.strip()])}\n"
        summary += f"- Total tabel: {table_count}\n"
        summary += f"- Total gambar: {image_count}\n"
        
        extracted_content.append(summary)
        
        result = "\n".join(extracted_content)
        logger.info(f"Advanced DOCX extraction completed: {len(result)} characters")
        return result
        
    except Exception as e:
        logger.error(f"Error dalam ekstraksi DOCX advanced: {e}")
        # Fallback ke docx2txt
        try:
            logger.info("Falling back to basic docx2txt extraction...")
            text = docx2txt.process(file_path)
            return f"=== DOKUMEN WORD (MODE SEDERHANA) ===\n{text}"
        except Exception as fallback_error:
            logger.error(f"Fallback juga gagal: {fallback_error}")
            return f"Error saat membaca DOCX: {str(e)}"

async def extract_text_from_pdf_advanced(file_path: str) -> str:
    """Ekstrak teks dari file PDF dengan PyMuPDF untuk hasil yang lebih baik."""
    try:
        logger.info("Starting advanced PDF extraction with PyMuPDF...")
        
        # Buka PDF dengan PyMuPDF
        pdf_document = fitz.open(file_path)
        extracted_content = []
        total_images = 0
        
        # Header
        extracted_content.append("=== DOKUMEN PDF ===\n")
        
        # Ekstrak metadata
        metadata = pdf_document.metadata
        if metadata.get('title'):
            extracted_content.append(f"Judul: {metadata['title']}")
        if metadata.get('author'):
            extracted_content.append(f"Penulis: {metadata['author']}")
        if metadata.get('subject'):
            extracted_content.append(f"Subjek: {metadata['subject']}")
        
        extracted_content.append(f"Total halaman: {len(pdf_document)}\n")
        
        # Ekstrak teks per halaman
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Ekstrak teks
            text = page.get_text()
            if text.strip():
                extracted_content.append(f"\n--- HALAMAN {page_num + 1} ---")
                extracted_content.append(text.strip())
            
            # Hitung gambar di halaman ini
            image_list = page.get_images()
            if image_list:
                page_images = len(image_list)
                total_images += page_images
                extracted_content.append(f"[{page_images} gambar ditemukan di halaman {page_num + 1}]")
        
        # Tambahkan ringkasan
        summary = f"\n=== RINGKASAN PDF ===\n"
        summary += f"- Total halaman: {len(pdf_document)}\n"
        summary += f"- Total gambar: {total_images}\n"
        
        extracted_content.append(summary)
        
        pdf_document.close()
        
        result = "\n".join(extracted_content)
        logger.info(f"Advanced PDF extraction completed: {len(result)} characters")
        return result
        
    except Exception as e:
        logger.error(f"Error dalam ekstraksi PDF advanced: {e}")
        # Fallback ke PyPDF4
        try:
            logger.info("Falling back to PyPDF4 extraction...")
            return extract_text_from_pdf(file_path)
        except Exception as fallback_error:
            logger.error(f"Fallback PDF juga gagal: {fallback_error}")
            return f"Error saat membaca PDF: {str(e)}"

def extract_text_from_pdf(file_path: str) -> str:
    """Ekstrak teks dari file PDF."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            # Menggunakan PyPDF4 yang pure Python
            pdf_reader = PyPDF4.PdfFileReader(file)
            for page_num in range(pdf_reader.getNumPages()):
                page = pdf_reader.getPage(page_num)
                text += page.extractText() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error saat mengekstrak teks dari PDF: {e}")
        return f"Error saat membaca PDF: {str(e)}"

def extract_text_from_docx(file_path: str) -> str:
    """Ekstrak teks dari file DOCX (fungsi fallback sederhana)."""
    try:
        text = docx2txt.process(file_path)
        return f"=== DOKUMEN WORD (MODE SEDERHANA) ===\n{text}" if text else "Dokumen kosong atau tidak dapat dibaca."
    except Exception as e:
        logger.error(f"Error saat mengekstrak teks dari DOCX: {e}")
        return f"Error saat membaca DOCX: {str(e)}"

# Endpoint untuk mengunggah dan memproses dokumen
@app.post("/api/upload_document", response_model=DocumentResponse)
async def upload_document(file: UploadFile = File(...)):
    """Upload dan proses dokumen (PDF atau DOCX)."""
    logger.info(f"Menerima permintaan upload file: {file.filename}")
    
    try:
        # Validasi tipe file dengan deteksi yang lebih akurat
        if not file.filename:
            logger.error("Filename is empty")
            raise HTTPException(
                status_code=400, 
                detail="Filename tidak boleh kosong"
            )
            
        file_extension = os.path.splitext(file.filename)[1].lower()
        logger.info(f"File extension: {file_extension}")
        
        # Daftar ekstensi yang didukung dengan kemampuan baru
        supported_extensions = ['.pdf', '.docx']
        supported_description = {
            '.pdf': 'PDF dengan support gambar dan metadata',
            '.docx': 'Word Document dengan support tabel, gambar, dan formatting'
        }
        
        if file_extension not in supported_extensions:
            logger.error(f"Unsupported file format: {file_extension}")
            supported_list = ", ".join([f"{ext} ({supported_description[ext]})" for ext in supported_extensions])
            raise HTTPException(
                status_code=400,
                detail=f"Format file tidak didukung. Format yang didukung: {supported_list}"
            )
        
        # Generate ID unik untuk dokumen
        document_id = str(uuid.uuid4())
        file_path = os.path.join(UPLOAD_FOLDER, f"{document_id}{file_extension}")
        logger.info(f"Saving file to: {file_path}")
        
        # Pastikan direktori upload ada
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # Simpan file
        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            logger.info(f"File saved successfully: {file_path}")
        except Exception as save_error:
            logger.error(f"Error saving file: {save_error}", exc_info=True)
            raise HTTPException(
                status_code=500, 
                detail=f"Gagal menyimpan file: {str(save_error)}"
            )
        
        # Ekstrak teks dari dokumen dengan metode advanced
        logger.info("Extracting text from document with advanced method...")
        document_text = await extract_text_from_document(file_path)
        logger.info(f"Text extraction completed, length: {len(document_text)} chars")
        
        # Analisis tambahan struktur dokumen
        analysis_info = ""
        try:
            if file_extension == '.docx':
                analysis = await analyze_docx_structure(file_path)
                analysis_info = f"\n=== INFO DOKUMEN ===\nParagraf: {analysis.get('paragraphs', 0)}, Tabel: {analysis.get('tables', 0)}, Gambar: {analysis.get('images', 0)}, Heading: {analysis.get('headings', 0)}\n"
            elif file_extension == '.pdf':
                analysis = await analyze_pdf_structure(file_path)
                analysis_info = f"\n=== INFO DOKUMEN ===\nHalaman: {analysis.get('pages', 0)}, Gambar: {analysis.get('images', 0)}, Panjang teks: {analysis.get('text_length', 0)} karakter\n"
        except Exception as analysis_error:
            logger.warning(f"Analysis failed but text extraction succeeded: {analysis_error}")
        
        # Gabungkan teks dokumen dengan info analisis
        full_content = document_text + analysis_info

        response = DocumentResponse(
            document_id=document_id,
            content=full_content,  # Return full content with analysis
            filename=file.filename
        )
        
        logger.info(f"Document processed successfully: {document_id}")
        
        # Add document metadata to library
        document_metadata = DocumentMetadata(
            document_id=document_id,
            filename=f"{document_id}{file_extension}",
            original_filename=file.filename,
            upload_date=datetime.now().isoformat(),
            file_size=os.path.getsize(file_path),
            file_type=file_extension,
            content_preview=full_content[:200],  # First 200 chars as preview
            analysis_summary={},  # Empty summary for now
            is_active=True  # Set as active document
        )
        
        document_library.add_document(document_metadata)
        
        # CRITICAL: Clear cache when uploading new document to prevent old context
        response_cache.clear()
        logger.info(f"🧹 Cache cleared after uploading new document: {file.filename}")
        
        # Enhanced response with chat context info
        return {
            "document_id": document_id,
            "content": full_content,
            "filename": file.filename,
            "chat_notification": {
                "type": "document_upload",
                "message": f"📄 **Document uploaded:** {file.filename}",
                "document_info": {
                    "document_id": document_id,
                    "filename": file.filename,
                    "file_type": file_extension,
                    "file_size": os.path.getsize(file_path),
                    "upload_date": datetime.now().isoformat()
                },
                "analysis_summary": analysis_info,
                "timestamp": datetime.now().isoformat()
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error saat mengunggah dokumen: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error saat mengunggah dokumen: {str(e)}")

# Document Library API Endpoints

@app.get("/api/documents")
async def get_document_library():
    """Get all documents in the library"""
    try:
        documents = document_library.get_all_documents()
        
        # Convert to response format
        response_docs = []
        for doc in documents:
            response_docs.append({
                "document_id": doc.document_id,
                "filename": doc.original_filename,
                "upload_date": doc.upload_date,
                "file_size": doc.file_size,
                "file_type": doc.file_type,
                "content_preview": doc.content_preview,
                "is_active": doc.is_active,
                "analysis_summary": doc.analysis_summary
            })
        
        # Sort by upload date (newest first)
        response_docs.sort(key=lambda x: x["upload_date"], reverse=True)
        
        return {
            "documents": response_docs,
            "total_count": len(response_docs),
            "active_document": next((doc for doc in response_docs if doc["is_active"]), None)
        }
        
    except Exception as e:
        logger.error(f"Error getting document library: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting documents: {str(e)}")

@app.get("/api/documents/{document_id}")
async def get_document_details(document_id: str):
    """Get detailed information about a specific document"""
    try:
        doc = document_library.get_document(document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get full content
        file_path = os.path.join(UPLOAD_FOLDER, doc.filename)
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Document file not found")
        
        # Extract full content again (cached version)
        full_content = await extract_text_from_document(file_path)
        
        return {
            "document_id": doc.document_id,
            "filename": doc.original_filename,
            "upload_date": doc.upload_date,
            "file_size": doc.file_size,
            "file_type": doc.file_type,
            "content": full_content,
            "content_preview": doc.content_preview,
            "is_active": doc.is_active,
            "analysis_summary": doc.analysis_summary
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document details: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting document: {str(e)}")

@app.post("/api/documents/{document_id}/select")
async def select_document(document_id: str):
    """Select a document as the active document for chat"""
    try:
        success = document_library.set_active_document(document_id)
        
        if not success:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # CRITICAL: Clear cache when switching documents to prevent old context
        response_cache.clear()
        logger.info("🧹 Cache cleared after document switch")
        
        # Get the newly active document
        active_doc = document_library.get_active_document()
        
        return {
            "success": True,
            "message": f"Document '{active_doc.original_filename}' selected successfully",
            "active_document": {
                "document_id": active_doc.document_id,
                "filename": active_doc.original_filename,
                "file_type": active_doc.file_type,
                "upload_date": active_doc.upload_date,
                "content_preview": active_doc.content_preview
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error selecting document: {e}")
        raise HTTPException(status_code=500, detail=f"Error selecting document: {str(e)}")

@app.delete("/api/documents/{document_id}")
async def delete_document_from_library(document_id: str):
    """Delete a document from the library"""
    try:
        # Get document info before deletion
        doc = document_library.get_document(document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        success = document_library.delete_document(document_id)
        
        if not success:
            raise HTTPException(status_code=500, detail="Failed to delete document")
        
        return {
            "success": True,
            "message": f"Document '{doc.original_filename}' deleted successfully",
            "deleted_document_id": document_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@app.post("/api/documents/clear-selection")
async def clear_document_selection():
    """Clear active document selection"""
    try:
        documents = document_library.get_all_documents()
        
        # Set all documents as inactive
        for doc in documents:
            doc.is_active = False
        
        document_library.save_metadata(documents)
        
        # CRITICAL: Clear cache when clearing document selection
        response_cache.clear()
        logger.info("🧹 Cache cleared after clearing document selection")
        
        return {
            "success": True,
            "message": "Document selection cleared"
        }
        
    except Exception as e:
        logger.error(f"Error clearing document selection: {e}")
        raise HTTPException(status_code=500, detail=f"Error clearing selection: {str(e)}")

@app.get("/api/documents/active")
async def get_active_document():
    """Get the currently active document"""
    try:
        active_doc = document_library.get_active_document()
        
        if not active_doc:
            return {
                "active_document": None,
                "message": "No active document selected"
            }
        
        # Get full content
        file_path = os.path.join(UPLOAD_FOLDER, active_doc.filename)
        if not os.path.exists(file_path):
            return {
                "active_document": None,
                "message": "Active document file not found"
            }
        
        # Extract full content again (cached version)
        full_content = await extract_text_from_document(file_path)
        
        return {
            "active_document": {
                "document_id": active_doc.document_id,
                "filename": active_doc.original_filename,
                "file_type": active_doc.file_type,
                "upload_date": active_doc.upload_date,
                "content_preview": active_doc.content_preview,
                "content": full_content
            }
        }
        
    except Exception as e:
        logger.error(f"Error getting active document: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting active document: {str(e)}")

# Document Context Change API
@app.post("/api/documents/{document_id}/context-switch")
async def document_context_switch(document_id: str):
    """Handle document context switch for chat notifications"""
    try:
        # Get document info
        doc = document_library.get_document(document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Set as active document
        success = document_library.set_active_document(document_id)
        if not success:
            raise HTTPException(status_code=500, detail="Failed to set active document")
        
        return {
            "success": True,
            "message": f"Context switched to: {doc.original_filename}",
            "chat_message": {
                "type": "system",
                "content": f"🔄 **Document context switched to:** {doc.original_filename}",
                "document_info": {
                    "document_id": doc.document_id,
                    "filename": doc.original_filename,
                    "file_type": doc.file_type,
                    "upload_date": doc.upload_date
                },
                "timestamp": datetime.now().isoformat()
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error switching document context: {e}")
        raise HTTPException(status_code=500, detail=f"Error switching context: {str(e)}")

# ═══════════════════════════════════════════════════════════════
# 📊 MULTI-DOCUMENT ANALYSIS ENDPOINTS
# ═══════════════════════════════════════════════════════════════

class MultiDocumentAnalysisRequest(BaseModel):
    message: str
    document_ids: List[str]
    mode: str = "sequential"  # sequential or batch

class MultiDocumentResult(BaseModel):
    document_id: str
    filename: str
    result: str
    status: str  # processing, completed, error
    processing_time: Optional[float] = None
    error_message: Optional[str] = None
    timestamp: datetime

class ProcessingProgress(BaseModel):
    total: int
    completed: int
    currentDocument: Optional[str] = None
    results: List[MultiDocumentResult]

class MultiDocumentAnalysisResponse(BaseModel):
    analysis_id: str
    status: str  # started, processing, completed, error
    results: List[MultiDocumentResult]
    progress: ProcessingProgress

# In-memory storage for multi-document analysis sessions
multi_doc_sessions: Dict[str, MultiDocumentAnalysisResponse] = {}

@app.post("/api/multi-document-analysis")
async def start_multi_document_analysis(request: MultiDocumentAnalysisRequest):
    """
    🚀 Start multi-document analysis process
    Processes multiple documents sequentially for optimal AI performance
    """
    try:
        # Generate analysis session ID
        analysis_id = str(uuid.uuid4())
        
        # Validate document IDs
        valid_documents = []
        for doc_id in request.document_ids:
            doc = document_library.get_document(doc_id)
            if doc:
                valid_documents.append(doc)
            else:
                logger.warning(f"Document {doc_id} not found in library")
        
        if not valid_documents:
            raise HTTPException(status_code=400, detail="No valid documents found")
        
        # Initialize analysis session
        analysis_session = MultiDocumentAnalysisResponse(
            analysis_id=analysis_id,
            status="started",
            results=[],
            progress=ProcessingProgress(
                total=len(valid_documents),
                completed=0,
                currentDocument=None,
                results=[]
            )
        )
        
        multi_doc_sessions[analysis_id] = analysis_session
        
        # Start background processing (in a real implementation, use async tasks)
        await process_multi_documents_sequential(
            analysis_id, 
            request.message, 
            valid_documents
        )
        
        return {
            "analysis_id": analysis_id,
            "status": "started",
            "message": f"Started analysis of {len(valid_documents)} documents"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error starting multi-document analysis: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis error: {str(e)}")

async def process_multi_documents_sequential(analysis_id: str, user_message: str, documents: List[DocumentMetadata]):
    """
    🔄 Process documents sequentially for optimal AI performance
    Updates session status in real-time
    """
    try:
        session = multi_doc_sessions[analysis_id]
        session.status = "processing"
        
        results = []
        
        for i, document in enumerate(documents):
            # Update current processing status
            session.progress.currentDocument = document.original_filename
            session.progress.completed = i
            
            try:
                start_time = time.time()
                
                # Get full document content
                file_path = os.path.join(UPLOAD_FOLDER, document.filename)
                if not os.path.exists(file_path):
                    raise Exception(f"Document file not found: {document.filename}")
                
                full_content = await extract_text_from_document(file_path)
                
                # Prepare context for this specific document
                document_context = f"""Document: {document.original_filename}
File Type: {document.file_type}
Content: {full_content[:8000]}"""  # Limit to 8KB for llama3:8b
                
                # Create optimized prompt for this document
                optimized_prompt = prompt_engineer.create_document_prompt(
                    user_message,
                    document_context,
                    []  # Fresh context for each document
                )
                
                # Send to AI using optimized configuration
                request_payload = ai_optimizer.get_optimized_payload(optimized_prompt)
                
                async with httpx.AsyncClient(timeout=300.0) as client:  # 5 minute timeout per document
                    response = await client.post(
                        "http://localhost:11434/api/generate",
                        json=request_payload,
                        timeout=300.0
                    )
                    
                    if response.status_code == 200:
                        result_data = response.json()
                        ai_response = result_data.get("response", "No response generated")
                        
                        # Clean the response
                        cleaned_response = response_monitor.clean_response(ai_response)
                        
                        processing_time = time.time() - start_time
                        
                        # Create result
                        result = MultiDocumentResult(
                            document_id=document.document_id,
                            filename=document.original_filename,
                            result=cleaned_response,
                            status="completed",
                            processing_time=round(processing_time, 2),
                            timestamp=datetime.now()
                        )
                        
                        results.append(result)
                        session.progress.results = results
                        session.progress.completed = i + 1
                        
                        logger.info(f"✅ Completed analysis for {document.original_filename} in {processing_time:.2f}s")
                    else:
                        raise Exception(f"Ollama API returned status {response.status_code}")
                
            except Exception as e:
                # Handle individual document processing error
                error_result = MultiDocumentResult(
                    document_id=document.document_id,
                    filename=document.original_filename,
                    result="",
                    status="error",
                    error_message=str(e),
                    timestamp=datetime.now()
                )
                
                results.append(error_result)
                session.progress.results = results
                session.progress.completed = i + 1
                
                logger.error(f"❌ Error processing {document.original_filename}: {e}")
        
        # Mark session as completed
        session.status = "completed"
        session.results = results
        session.progress.currentDocument = None
        
        logger.info(f"🎉 Multi-document analysis {analysis_id} completed: {len(results)} documents processed")
        
    except Exception as e:
        # Mark session as error
        session.status = "error"
        logger.error(f"💥 Fatal error in multi-document analysis {analysis_id}: {e}")

@app.get("/api/multi-document-analysis/{analysis_id}/status")
async def get_analysis_status(analysis_id: str):
    """
    📊 Get real-time status of multi-document analysis
    Used for polling progress updates
    """
    try:
        if analysis_id not in multi_doc_sessions:
            raise HTTPException(status_code=404, detail="Analysis session not found")
        
        session = multi_doc_sessions[analysis_id]
        
        return {
            "analysis_id": analysis_id,
            "status": session.status,
            "progress": {
                "total": session.progress.total,
                "completed": session.progress.completed,
                "currentDocument": session.progress.currentDocument,
                "results": [
                    {
                        "document_id": r.document_id,
                        "filename": r.filename,
                        "result": r.result,
                        "status": r.status,
                        "processing_time": r.processing_time,
                        "error_message": r.error_message,
                        "timestamp": r.timestamp.isoformat()
                    } for r in session.progress.results
                ]
            },
            "results": [
                {
                    "document_id": r.document_id,
                    "filename": r.filename,
                    "result": r.result,
                    "status": r.status,
                    "processing_time": r.processing_time,
                    "error_message": r.error_message,
                    "timestamp": r.timestamp.isoformat()
                } for r in session.results
            ]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting analysis status: {e}")
        raise HTTPException(status_code=500, detail=f"Status error: {str(e)}")

# Run server directly if this file is executed
if __name__ == "__main__":
    import uvicorn
    
    print("🚀 Starting Dokai Chat Backend Server...")
    print("📡 Server will be available at: http://localhost:8000")
    print("📚 API Documentation: http://localhost:8000/docs")
    print("🔄 Auto-reload enabled for development")
    
    try:
        uvicorn.run(
            "main:app",
            host="0.0.0.0",
            port=8000,
            reload=True,
            log_level="info"
        )
    except KeyboardInterrupt:
        print("\n🛑 Server stopped by user")
    except Exception as e:
        print(f"❌ Error starting server: {e}")
        import traceback
        traceback.print_exc()
